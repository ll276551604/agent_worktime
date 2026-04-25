# -*- coding: utf-8 -*-
import os
import sys
import uuid
import json
import queue
import threading
import time
import logging
from logging.handlers import RotatingFileHandler
from datetime import datetime

from flask import Flask, request, jsonify, Response, send_file, render_template
from werkzeug.utils import secure_filename

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from config import UPLOAD_FOLDER, OUTPUT_FOLDER, MAX_UPLOAD_SIZE, AVAILABLE_MODELS, DEFAULT_MODEL, AppConfig


def setup_logging():
    """配置日志系统"""
    logger = logging.getLogger()
    logger.setLevel(logging.INFO)
    
    # 避免重复添加处理器
    if logger.handlers:
        logger.handlers.clear()
    
    # 格式设置
    formatter = logging.Formatter(
        '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    
    # 控制台输出
    console_handler = logging.StreamHandler()
    console_handler.setFormatter(formatter)
    logger.addHandler(console_handler)
    
    # 文件输出（自动轮转，保留最近5个文件，每个最大10MB）
    file_handler = RotatingFileHandler(
        'app.log',
        maxBytes=10*1024*1024,  # 10MB
        backupCount=5,
        encoding='utf-8'
    )
    file_handler.setFormatter(formatter)
    logger.addHandler(file_handler)


# 初始化日志
setup_logging()
logger = logging.getLogger(__name__)

# 初始化目录
AppConfig.init_folders()

from excel import reader
from agent import worktime_agent

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = MAX_UPLOAD_SIZE
app.config['TIMEOUT'] = 300  # 5分钟超时


# ============================================================
# 结构化追问（LLM 驱动）
# ============================================================
def _generate_clarifying_question(conversation_history: list, dm, model_id: str, existing_requirement: str = "") -> str:
    """使用 LLM 生成结构化的澄清追问，帮助用户完善需求描述"""
    from agent import gemini_client

    # 构建上下文
    ctx_lines = []
    for msg in conversation_history[-6:]:  # 最近6条消息
        role = "用户" if msg["role"] == "user" else "助手"
        ctx_lines.append(f"{role}: {msg['content']}")
    context = "\n".join(ctx_lines) if ctx_lines else "（无历史对话）"

    req_line = f"\n当前需求描述: {existing_requirement}" if existing_requirement else ""

    prompt = f"""你是一位专业的产品经理，正在引导用户描述产品需求。

## 对话上下文
{context}{req_line}

## 任务
用户的需求描述不够完整，请生成 1~2 个结构化的追问问题，帮助用户补充关键信息。

## 追问原则
1. 先判断用户已经说了什么，再判断缺什么
2. 每次最多问 2 个问题，问题要具体、有针对性
3. 问题格式：用自然口语化的方式提问，末尾加「~」
4. 如果用户已经描述了功能，追问应聚焦：涉及哪些页面/模块？是新增还是改造？涉及哪些角色？
5. 如果用户只说了模块名，追问：具体要做什么功能？是新增还是调整？

## 输出格式
直接输出追问内容，不要任何其他说明。

示例1（需求过短）：
根据您的描述「店铺管理新增字段」，我想再确认几个细节~
1. 这个字段在哪些页面需要展示？（列表页、详情页、还是其他页面？）
2. 只是展示还是需要支持编辑？

示例2（缺少类型）：
我理解您想在「订单管理」模块做调整~ 请问具体是：
1. 新增什么功能？还是对已有功能做优化？
2. 涉及哪些页面和操作流程？
"""

    try:
        response = gemini_client.call_llm(prompt, model_id=model_id)
        return response.strip()
    except Exception:
        # 降级为规则追问
        return dm.generate_intelligent_question({"requirement": existing_requirement, "module": "", "type": ""})


# ============================================================
# 统一响应格式
# ============================================================
def make_response(output_type, content, intent=None, stage=None, session_id=None, meta=None):
    """
    统一响应格式生成器
    :param output_type: 输出类型: text | evaluation | thinking | error
    :param content: 内容
    :param intent: 意图
    :param stage: 阶段
    :param session_id: 会话ID
    :param meta: 额外元数据
    """
    response = {
        "success": output_type != "error",
        "output": {
            "type": output_type,
            "content": content
        },
        "meta": {
            "intent": intent,
            "stage": stage,
            "session_id": session_id
        }
    }
    if meta:
        response["meta"].update(meta)
    return jsonify(response)


# ============================================================
# 全局异常处理
# ============================================================
@app.errorhandler(Exception)
def handle_exception(e):
    """全局异常处理器"""
    logger.error(f"未处理的异常: {str(e)}", exc_info=True)
    return jsonify({
        "success": False,
        "error": "服务器内部错误，请稍后重试"
    }), 500


@app.errorhandler(400)
def handle_bad_request(e):
    """处理请求参数错误"""
    return jsonify({
        "success": False,
        "error": str(e)
    }), 400


@app.errorhandler(403)
def handle_forbidden(e):
    """处理禁止访问"""
    return jsonify({
        "success": False,
        "error": "访问被拒绝"
    }), 403


@app.errorhandler(404)
def handle_not_found(e):
    """处理资源未找到"""
    return jsonify({
        "success": False,
        "error": "资源未找到"
    }), 404


@app.errorhandler(413)
def handle_request_too_large(e):
    """处理文件过大"""
    return jsonify({
        "success": False,
        "error": f"上传文件大小超过限制（最大 {MAX_UPLOAD_SIZE//1024//1024}MB）"
    }), 413



@app.route("/")
def index():
    return render_template("index.html")


@app.route("/models")
def models():
    return jsonify({
        "models":  [{"id": m["id"], "label": m["label"]} for m in AVAILABLE_MODELS],
        "default": DEFAULT_MODEL,
    })


# ============================================================
# 会话管理接口
# ============================================================
@app.route("/session/create", methods=["POST"])
def create_session():
    """创建新会话"""
    from agent.session_manager import SessionManager
    sm = SessionManager()
    session = sm.create_session()
    logger.info(f"创建会话: {session.session_id}")
    return jsonify({
        "session_id": session.session_id,
        "message": "会话创建成功",
    })


@app.route("/session/<session_id>/history")
def get_session_history(session_id):
    """获取会话历史"""
    from agent.session_manager import SessionManager
    sm = SessionManager()
    session = sm.get_session(session_id)
    if not session:
        return jsonify({"error": "会话不存在"}), 404
    return jsonify({
        "session_id": session_id,
        "history": session.get_history(),
    })


def _parse_document(file_bytes, ext, filename):
    """解析不同格式的文档"""
    try:
        if ext == ".pdf":
            try:
                import PyPDF2
                from io import BytesIO
                pdf_file = BytesIO(file_bytes)
                reader = PyPDF2.PdfReader(pdf_file)
                text_parts = []
                for page in reader.pages:
                    text_parts.append(page.extract_text())
                return "\n".join(text_parts)
            except:
                return None
        elif ext == ".xlsx":
            try:
                from openpyxl import load_workbook
                from io import BytesIO
                excel_file = BytesIO(file_bytes)
                wb = load_workbook(excel_file)
                text_parts = []
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    text_parts.append(f"\n=== {sheet_name} ===\n")
                    for idx, row in enumerate(ws.iter_rows(values_only=True)):
                        if idx > 100:
                            text_parts.append("... [表格过长，已截断]")
                            break
                        row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                        text_parts.append(row_text)
                return "\n".join(text_parts)
            except:
                return None
        elif ext == ".docx":
            try:
                from docx import Document
                from io import BytesIO
                docx_file = BytesIO(file_bytes)
                doc = Document(docx_file)
                text_parts = []
                for para in doc.paragraphs:
                    text_parts.append(para.text)
                for table in doc.tables:
                    text_parts.append("\n[表格]\n")
                    for row in table.rows:
                        row_text = " | ".join(cell.text for cell in row.cells)
                        text_parts.append(row_text)
                return "\n".join(text_parts)
            except:
                return None
        elif ext in [".jpg", ".jpeg", ".png", ".gif", ".bmp"]:
            return f"[图片文件: {ext}]\n(暂不支持自动OCR，请使用PDF或其他文本格式)"
        elif ext in [".txt", ".md", ".log"]:
            return file_bytes.decode('utf-8', errors='ignore')
        else:
            return None
    except Exception as e:
        logger.error(f"文档解析失败 ({filename}): {e}")
        return None


def _extract_interfaces(doc_content):
    """从文档内容中提取接口名称"""
    import re
    interfaces = set()
    patterns = [
        r'(?:接口|API|api|endpoint)[\s：:]*([a-zA-Z0-9_\.\-/]+)',
        r'(?:方法|Method)[\s：:]*(?:GET|POST|PUT|DELETE)\s+(/[a-zA-Z0-9_/\-]*)',
        r'def\s+([a-zA-Z0-9_]+)\s*\(',
        r'function\s+([a-zA-Z0-9_]+)\s*\(',
        r'public\s+(?:void|String|int|boolean)\s+([a-zA-Z0-9_]+)\s*\(',
    ]
    for pattern in patterns:
        matches = re.findall(pattern, doc_content)
        interfaces.update(matches)
    return list(interfaces)[:20]


@app.route("/upload_knowledge", methods=["POST"])
def upload_knowledge():
    """用户上传知识库文档（用于当前对话的上下文）"""
    if "file" not in request.files:
        return jsonify({"error": "未收到文件"}), 400

    # 从query参数或form获取session_id
    session_id = request.args.get("session_id") or request.form.get("session_id")

    f = request.files["file"]
    if not f.filename:
        return jsonify({"error": "文件名为空"}), 400

    try:
        filename = secure_filename(f.filename)
        file_bytes = f.read()
        ext = os.path.splitext(filename)[1].lower()

        doc_content = _parse_document(file_bytes, ext, filename)
        if not doc_content:
            return jsonify({"error": "无法解析该文件格式"}), 400

        max_content_length = 10000
        if len(doc_content) > max_content_length:
            doc_content = doc_content[:max_content_length] + "\n... [文档过长，已截断]"

        interfaces = _extract_interfaces(doc_content)

        from agent.session_manager import SessionManager
        session_mgr = SessionManager()

        # 如果没有session_id或session不存在，创建新session
        session = None
        if session_id:
            session = session_mgr.get_session(session_id)

        if not session:
            session = session_mgr.create_session()
            session_id = session.session_id

        session.add_temp_document({
            "filename": filename,
            "content": doc_content,
            "interfaces": interfaces,
            "upload_time": datetime.now().isoformat()
        })

        logger.info(f"临时文档上传: session={session_id}, file={filename}, size={len(file_bytes)}bytes")

        return jsonify({
            "success": True,
            "session_id": session_id,
            "filename": filename,
            "document_count": len(session.temp_documents),
            "interfaces": interfaces[:10],
            "message": f"已上传 {filename}，提取到 {len(interfaces)} 个接口"
        })

    except Exception as e:
        logger.error(f"文档上传失败: {e}")
        return jsonify({"error": f"处理失败: {str(e)}"}), 500


@app.route("/session/<session_id>/delete", methods=["POST"])
def delete_session(session_id):
    """删除会话"""
    from agent.session_manager import SessionManager
    sm = SessionManager()
    sm.delete_session(session_id)
    return jsonify({"message": "会话已删除"})


@app.route("/chat", methods=["POST"])
def chat():
    """聊天接口 — 产品需求拆解&工时评估专用工作流Agent"""
    from agent.session_manager import SessionManager
    from agent.dialog_manager import DialogManager
    from agent import skill_manager as sm
    import json

    session_mgr = SessionManager()
    dm          = DialogManager()
    data        = request.get_json()

    session_id       = data.get("session_id")
    message         = (data.get("message") or "").strip()
    model_id        = data.get("model_id", DEFAULT_MODEL)
    skill_id        = data.get("skill_id") or sm.get_current_skill_id()

    if not message:
        return jsonify({"error": "消息内容为空"}), 400

    # ── 获取/创建会话 ────────────────────────────────────────
    if session_id:
        session = session_mgr.get_session(session_id)
    if not session_id or not session:
        session = session_mgr.create_session()
        session_id = session.session_id

    # 获取历史对话和评估结果（用于意图识别）
    conversation_history = session.get_messages()
    last_evaluation = session.get_last_evaluation()
    has_history = len(conversation_history) > 0
    has_evaluation = bool(last_evaluation)

    # ── 强制第一步：全局意图识别（所有用户输入，永远优先执行） ──
    intent_result = dm.analyze_intent(message, has_history, has_evaluation)
    intent_data = json.loads(intent_result)
    intent = intent_data["intent"]
    intent_reason = intent_data["reason"]
    
    logger.info(f"意图识别结果: session={session_id} intent={intent} reason={intent_reason}")

    # 添加用户消息到会话
    session.add_message("user", message)
    conversation_history = session.get_messages()

    # ── 1. intent = chat（闲聊） ────────────────────────────
    if intent == "chat":
        # 纠正性闲聊回复 - 简洁引导用户回到业务主题
        chat_responses = {
            "你好": "您好！我是产品需求拆解&工时评估专用Agent 😊\n\n请直接描述您的产品需求，我会帮您拆解并评估工时。\n\n示例：开发一个用户登录功能",
            "hello": "Hello! I'm a dedicated Agent for product requirement breakdown and worktime estimation. 😊\n\nPlease describe your product requirement directly.\n\nExample: Develop a user login feature",
            "hi": "Hi! 😊 我是产品需求拆解&工时评估专用Agent。\n\n请直接描述您的产品需求，我会帮您拆解评估。\n\n示例：开发一个订单管理模块",
            "您好": "您好！我是产品需求拆解&工时评估专用Agent 😊\n\n请直接描述您的产品需求，我会帮您拆解评估。\n\n示例：实现报表系统的数据可视化功能",
            "早上好": "早上好！🌞\n\n我是产品需求拆解&工时评估专用Agent。请直接描述您的产品需求。\n\n示例：开发用户注册功能",
            "晚上好": "晚上好！🌙\n\n我是产品需求拆解&工时评估专用Agent。请直接描述您的产品需求。\n\n示例：优化订单列表页面性能",
            "下午好": "下午好！😊\n\n我是产品需求拆解&工时评估专用Agent。请直接描述您的产品需求。\n\n示例：开发商品管理后台",
            "谢谢": "不客气！😊\n\n如果您有产品需求拆解或工时评估的需求，随时可以找我。",
            "谢谢了": "不客气！😊\n\n如果您有产品需求需要评估，随时可以回来找我。",
            "感谢": "不客气！很高兴能帮到您。😊",
            "拜拜": "再见！👋\n\n如果您有产品需求拆解或工时评估的需求，欢迎随时回来。",
            "再见": "再见！👋\n\n如果您有产品需求需要评估，欢迎随时回来。",
            "你是谁": "我是产品需求拆解&工时评估专用Agent！🤖\n\n我的核心能力：\n• 需求分析与拆解\n• 工时评估与核算\n• 支持多轮迭代调整\n\n请直接描述您的产品需求，我会帮您评估！",
            "你叫什么": "我是产品需求拆解&工时评估专用Agent！😊\n\n请直接描述您的产品需求，我会帮您拆解评估。",
            "你能做什么": "我可以帮您：🤝\n\n1. **需求拆解**：将产品需求拆分成具体功能模块\n2. **工时评估**：套用官方工时评估标准核算工时\n3. **多轮迭代**：支持补充订正后重新评估\n\n请直接描述您的产品需求！",
            "你的功能": "我的主要功能是产品需求拆解和工时评估！📊\n\n核心能力：\n• 智能识别需求类型\n• 全量结构化拆解\n• 工时核算评估\n• 支持多轮调整\n\n请直接描述您的产品需求！",
            "介绍一下": "我是产品需求拆解&工时评估专用Agent！🤖\n\n我能帮您拆解产品需求并评估工时，支持多轮迭代调整。\n\n请直接描述您的产品需求！",
            "规则": "我会按照以下规则评估：📋\n\n1. 已有功能=调整需求，新能力=新增需求\n2. 按模块、页面、功能点结构化拆解\n3. 套用官方工时评估标尺核算\n4. 支持补充纠错后全流程重算\n\n请直接描述您的产品需求！",
            "说明": "我会：📝\n\n1. 识别需求类型（新增/调整/混合）\n2. 全量结构化拆解需求\n3. 核算工时评估\n\n请直接描述您的产品需求！",
            "帮助": "当然可以！🙋\n\n您可以直接描述产品需求，我会帮您拆解和评估。\n\n如果需要调整，随时补充说明即可重新评估。",
            "使用方法": "使用方法：💡\n\n1. 输入需求 → 2. 查看评估 → 3. 调整优化\n\n请直接描述您的产品需求！\n\n示例：开发一个用户登录功能",
            "怎么用": "直接输入产品需求即可！💡\n\n示例：\n• 开发订单管理模块\n• 修复首页加载缓慢问题\n• 优化用户注册流程",
            "什么是": "📚\n\n**需求拆解**：将产品需求拆分成具体功能模块、页面、接口等\n**工时评估**：根据拆解结果核算开发所需时间\n\n请直接描述您的产品需求，我会帮您评估！",
        }
        
        # 查找匹配的回复，如果没有匹配则给出通用回复
        base_reply = chat_responses.get(message, None)
        if base_reply:
            reply = base_reply
        else:
            # 对于未匹配的闲聊内容，给出纠正性引导
            reply = f"不好意思，我是产品需求拆解&工时评估专用Agent，主要专注于产品需求拆解和工时评估工作。😊\n\n温馨提示：如果您有产品需求需要评估，请直接描述您的需求，例如：\n「开发一个用户登录功能」\n\n请问有什么需求需要我帮忙拆解评估吗？"
        
        session.add_message("assistant", reply)
        return make_response(
            output_type="text",
            content=reply,
            intent=intent,
            stage="chat",
            session_id=session_id,
            meta={
                "intent_reason": intent_reason,
                "process": [f"意图识别：chat - {intent_reason}"]
            }
        )

    # ── 2. intent = new_task（全新需求拆解任务） ──────────────
    # ── 3. intent = revise_task（订正重跑任务） ──────────────
    # 两者共享相同的评估流程，revise_task 会基于历史上下文重新评估
    if intent in ["new_task", "revise_task"]:
        process_log = [f"意图识别：{intent} - {intent_reason}"]
        
        # ── 提取需求信息（模块/类型自动识别，无需用户确认） ────────────────
        info = dm.extract_requirement_info(conversation_history)
        auto_detected = {
            "module_detected": bool(info.get("auto_detected", {}).get("module")),
            "type_detected": bool(info.get("auto_detected", {}).get("type")),
        }

        # ── 检查是否有有效需求描述 ────────────────
        requirement = info.get('requirement', '').strip()

        if not requirement:
            # 没有有效需求，使用 LLM 生成结构化追问
            process_log.append("未识别到有效需求描述，生成结构化追问")
            question = _generate_clarifying_question(conversation_history, dm, model_id)
            session.add_message("assistant", question)
            return make_response(
                output_type="text",
                content=question,
                intent=intent,
                stage="clarifying",
                session_id=session_id,
                meta={
                    "intent_reason": intent_reason,
                    "collected_info": info,
                    "auto_detected": auto_detected,
                    "process": process_log,
                }
            )

        # 需求描述过短，使用 LLM 生成结构化追问
        if len(requirement) < 15:
            process_log.append("需求描述过短，生成结构化追问")
            question = _generate_clarifying_question(conversation_history, dm, model_id, requirement)
            session.add_message("assistant", question)
            return make_response(
                output_type="text",
                content=question,
                intent=intent,
                stage="clarifying",
                session_id=session_id,
                meta={
                    "intent_reason": intent_reason,
                    "collected_info": info,
                    "auto_detected": auto_detected,
                    "process": process_log,
                }
            )

        process_log.append(f"自动识别 → 模块: {info.get('module', '未指定')} 类型: {info.get('type', '新增功能')}")

        # ── 拼装完整需求文本（含模块+类型上下文） ────────────
        full_text = f"【模块】{info['module']} 【类型】{info['type']} 【描述】{info['requirement']}"
        process_log.append(f"提取到需求：{info['requirement'][:80]}")
        if info.get('module'):
            process_log.append(f"模块识别：{info['module']} ({'自动' if auto_detected['module_detected'] else '默认'})")
        else:
            process_log.append("模块未识别，使用默认模块")
        if info.get('type'):
            process_log.append(f"需求类型识别：{info['type']} ({'自动' if auto_detected['type_detected'] else '默认'})")
        else:
            process_log.append("需求类型未识别，默认视为新增功能")

        # ── 获取完整历史对话上下文（用于 revise_task 全流程重跑） ──
        # 过滤闲聊消息，只保留需求和评估相关内容
        chat_keywords = ['你好', 'hello', 'hi', '您好', '早上好', '晚上好', '下午好', '嗨',
                        '谢谢', '谢谢了', '感谢', '拜拜', '再见', '再见了', '回见',
                        '你是谁', '你叫什么', '你能做什么', '你的功能', '介绍一下',
                        '规则', '说明', '帮助', '使用方法', '怎么用', '什么是']
        
        msgs_ctx = []
        for m in conversation_history[:-1]:  # 排除刚加入的本条消息
            content = m['content'].lower()
            if not any(keyword in content for keyword in chat_keywords):
                msgs_ctx.append(m)
        
        # 限制上下文长度（3000字符），优先保留最新消息
        max_total_length = 3000
        history_ctx = ""
        for m in reversed(msgs_ctx):
            msg_text = f"{m['role']}: {m['content']}"
            if len(history_ctx) + len(msg_text) <= max_total_length:
                history_ctx = msg_text + "\n" + history_ctx
            else:
                remaining = max_total_length - len(history_ctx) - len(m['role']) - 2
                if remaining > 0:
                    history_ctx = f"{m['role']}: {m['content'][:remaining]}...\n" + history_ctx
                break
        
        process_log.append(f"历史上下文长度: {len(history_ctx)}")

        if intent == "revise_task":
            process_log.append("执行全流程重跑：重新判定需求类型、重新全量结构化拆解、重新完整工时评估")
        else:
            process_log.append("执行全新需求拆解评估流程")

        # ── 执行需求拆解与工时评估 ────────────────────────────
        session_knowledge = session.get_temp_knowledge_context()
        try:
            result = worktime_agent.run_chat(
                text=full_text,
                model_id=model_id,
                context=history_ctx,
                skill_id=skill_id,
                last_evaluation=last_evaluation,
                session_knowledge=session_knowledge,
            )
        except Exception as e:
            logger.error(f"需求拆解评估失败: {e}", exc_info=True)
            return jsonify({"error": f"评估失败: {str(e)}"}), 500

        # ── 需要澄清：需求描述过于简短 ────────────────────────
        if result.get("needs_clarification"):
            process_log.append("评估结果认为需求描述过短，需要补充信息")
            question = result["clarification_question"]
            session.add_message("assistant", question)
            return make_response(
                output_type="text",
                content=question,
                intent=intent,
                stage="clarifying",
                session_id=session_id,
                meta={
                    "intent_reason": intent_reason,
                    "collected_info": info,
                    "auto_detected": auto_detected,
                    "process": process_log,
                }
            )

    # ── 追问模式（已有上下文，用户在追问） ──────────────────
    if result.get("is_question"):
        process_log.append("检测到用户追问，直接生成回答")
        session.add_message("assistant", result["g_text"])
        return make_response(
            output_type="text",
            content=result["g_text"],
            intent=intent,
            stage="answering",
            session_id=session_id,
            meta={
                "intent_reason": intent_reason,
                "collected_info": info,
                "process": process_log,
            }
        )

    # ── 反馈重新评估结果 ───────────────────────────────────────
    if result.get("is_feedback"):
        process_log.append("检测到用户反馈，执行重新评估")
        session.add_message("assistant", result["g_text"])
        logger.info(f"反馈重新评估完成: session={session_id} days={result['total_days']}")
        
        # 更新上次评估结果
        session.set_last_evaluation({
            "total_days": result["total_days"],
            "role_breakdown": result.get("role_breakdown", {}),
            "g_text": result["g_text"],
        })
        
        return make_response(
            output_type="evaluation",
            content={
                "g_text": result["g_text"],
                "total_days": result["total_days"],
                "role_breakdown": result.get("role_breakdown", {}),
                "pages_features": [],
            },
            intent=intent,
            stage="reassessment",
            session_id=session_id,
            meta={
                "intent_reason": intent_reason,
                "skill_id": skill_id,
                "collected_info": info,
                "auto_detected": auto_detected,
                "process": process_log,
                "is_feedback": True,
            }
        )

    # ── 正常评估结果 ─────────────────────────────────────────
    pages_features = result.get("pages_features", [])
    role_breakdown = result.get("role_breakdown", {})

    # 生成表格格式的输出
    table_output = worktime_agent.format_evaluation_as_table({
        "pages_features": pages_features,
        "total_days": result["total_days"],
        "role_breakdown": role_breakdown,
    }, session_knowledge=session_knowledge)

    # 智能识别提示头 + 表格格式
    formatted_result = table_output
    note_parts = []
    if auto_detected["module_detected"]:
        note_parts.append(f"模块：{info['module']}")
    if auto_detected["type_detected"]:
        note_parts.append(f"类型：{info['type']}")
    if note_parts:
        formatted_result = f"【智能识别】{'、'.join(note_parts)}\n\n" + formatted_result

    process_log.append("评估完成，生成表格格式的拆解与工时结果")
    session.add_message("assistant", formatted_result)
    logger.info(f"评估完成: session={session_id} skill={skill_id} "
                f"pages={len(pages_features)} days={result['total_days']}")
    
    # 保存评估结果到会话（用于后续反馈重新评估）
    session.set_last_evaluation({
        "total_days": result["total_days"],
        "role_breakdown": role_breakdown,
        "g_text": result["g_text"],
        "pages_features": pages_features,
    })

    return make_response(
        output_type="evaluation",
        content={
            "g_text": formatted_result,
            "total_days": result["total_days"],
            "role_breakdown": role_breakdown,
            "pages_features": pages_features,
            # 向下兼容字段
            "decomposition": [
                {"type": p["类型"], "name": p["页面"], "features": p["功能点"]}
                for p in pages_features
            ],
            "evaluation": {
                "model": f"skill:{skill_id}",
                "effort_days": result["total_days"],
                "role_breakdown": role_breakdown,
            },
        },
        intent=intent,
        stage="assessment",
        session_id=session_id,
        meta={
            "intent_reason": intent_reason,
            "skill_id": skill_id,
            "collected_info": info,
            "auto_detected": auto_detected,
            "process": process_log,
        }
    )


@app.route("/chat/stream", methods=["POST"])
def chat_stream():
    """流式聊天接口 — 支持thinking状态输出和逐步加载"""
    from agent.session_manager import SessionManager
    from agent.dialog_manager import DialogManager
    from agent import skill_manager as sm
    import json
    import time

    session_mgr = SessionManager()
    dm = DialogManager()
    data = request.get_json()

    session_id = data.get("session_id")
    message = (data.get("message") or "").strip()
    model_id = data.get("model_id", DEFAULT_MODEL)
    skill_id = data.get("skill_id") or sm.get_current_skill_id()

    if not message:
        return jsonify({"error": "消息内容为空"}), 400

    # 获取/创建会话
    if session_id:
        session = session_mgr.get_session(session_id)
    if not session_id or not session:
        session = session_mgr.create_session()
        session_id = session.session_id

    # 获取历史对话和评估结果
    conversation_history = session.get_messages()
    last_evaluation = session.get_last_evaluation()
    has_history = len(conversation_history) > 0
    has_evaluation = bool(last_evaluation)

    # 意图识别
    intent_result = dm.analyze_intent(message, has_history, has_evaluation)
    intent_data = json.loads(intent_result)
    intent = intent_data["intent"]
    intent_reason = intent_data["reason"]

    session.add_message("user", message)
    conversation_history = session.get_messages()

    def generate():
        nonlocal intent, intent_reason, session, conversation_history, last_evaluation
        
        # 发送意图识别结果
        intent_data = {'type': 'intent', 'intent': intent, 'reason': intent_reason}
        yield 'data: ' + json.dumps(intent_data) + '\n\n'
        
        # 处理闲聊意图
        if intent == "chat":
            chat_responses = {
                "你好": "您好！我是产品需求拆解&工时评估专用Agent 😊\n\n请直接描述您的产品需求，我会帮您拆解并评估工时。\n\n示例：开发一个用户登录功能",
                "hello": "Hello! I'm a dedicated Agent for product requirement breakdown. 😊\n\nPlease describe your product requirement directly.\n\nExample: Develop a user login feature",
                "hi": "Hi! 😊 我是产品需求拆解&工时评估专用Agent。\n\n请直接描述您的产品需求，我会帮您拆解评估。\n\n示例：开发一个订单管理模块",
                "您好": "您好！我是产品需求拆解&工时评估专用Agent 😊\n\n请直接描述您的产品需求，我会帮您拆解评估。\n\n示例：实现报表系统的数据可视化功能",
                "早上好": "早上好！🌞\n\n我是产品需求拆解&工时评估专用Agent。请直接描述您的产品需求。\n\n示例：开发用户注册功能",
                "晚上好": "晚上好！🌙\n\n我是产品需求拆解&工时评估专用Agent。请直接描述您的产品需求。\n\n示例：优化订单列表页面性能",
                "下午好": "下午好！😊\n\n我是产品需求拆解&工时评估专用Agent。请直接描述您的产品需求。\n\n示例：开发商品管理后台",
                "谢谢": "不客气！😊\n\n如果您有产品需求拆解或工时评估的需求，随时可以找我。",
                "谢谢了": "不客气！😊\n\n如果您有产品需求需要评估，随时可以回来找我。",
                "感谢": "不客气！很高兴能帮到您。😊",
                "拜拜": "再见！👋\n\n如果您有产品需求拆解或工时评估的需求，欢迎随时回来。",
                "再见": "再见！👋\n\n如果您有产品需求需要评估，欢迎随时回来。",
                "你是谁": "我是产品需求拆解&工时评估专用Agent！🤖\n\n我的核心能力：\n• 需求分析与拆解\n• 工时评估与核算\n• 支持多轮迭代调整\n\n请直接描述您的产品需求！",
                "你叫什么": "我是产品需求拆解&工时评估专用Agent！😊\n\n请直接描述您的产品需求，我会帮您拆解评估。",
                "你能做什么": "我可以帮您：🤝\n\n1. **需求拆解**：将产品需求拆分成具体功能模块\n2. **工时评估**：套用官方工时评估标准核算工时\n3. **多轮迭代**：支持补充订正后重新评估\n\n请直接描述您的产品需求！",
                "你的功能": "我的主要功能是产品需求拆解和工时评估！📊\n\n核心能力：\n• 智能识别需求类型\n• 全量结构化拆解\n• 工时核算评估\n• 支持多轮调整\n\n请直接描述您的产品需求！",
                "介绍一下": "我是产品需求拆解&工时评估专用Agent！🤖\n\n我能帮您拆解产品需求并评估工时，支持多轮迭代调整。\n\n请直接描述您的产品需求！",
                "规则": "我会按照以下规则评估：📋\n\n1. 已有功能=调整需求，新能力=新增需求\n2. 按模块、页面、功能点结构化拆解\n3. 套用官方工时评估标尺核算\n4. 支持补充纠错后全流程重算\n\n请直接描述您的产品需求！",
                "说明": "我会：📝\n\n1. 识别需求类型（新增/调整/混合）\n2. 全量结构化拆解需求\n3. 核算工时评估\n\n请直接描述您的产品需求！",
                "帮助": "当然可以！🙋\n\n您可以直接描述产品需求，我会帮您拆解和评估。\n\n如果需要调整，随时补充说明即可重新评估。",
                "使用方法": "使用方法：💡\n\n1. 输入需求 → 2. 查看评估 → 3. 调整优化\n\n请直接描述您的产品需求！\n\n示例：开发一个用户登录功能",
                "怎么用": "直接输入产品需求即可！💡\n\n示例：\n• 开发订单管理模块\n• 修复首页加载缓慢问题\n• 优化用户注册流程",
                "什么是": "📚\n\n**需求拆解**：将产品需求拆分成具体功能模块、页面、接口等\n**工时评估**：根据拆解结果核算开发所需时间\n\n请直接描述您的产品需求，我会帮您评估！",
            }
            # 查找匹配的回复，如果没有匹配则给出通用回复
            base_reply = chat_responses.get(message, None)
            if base_reply:
                reply = base_reply
            else:
                # 对于未匹配的闲聊内容，给出简洁的纠正性引导
                reply = "我是产品需求拆解&工时评估专用Agent。请直接描述您的产品需求，我会帮您拆解评估。\n\n示例：开发一个用户登录功能"
            
            complete_data = {'type': 'complete', 'stage': 'chat', 'message': reply}
            yield 'data: ' + json.dumps(complete_data) + '\n\n'
            return
        
        # 处理评估任务
        if intent in ["new_task", "revise_task"]:
            # 提取需求信息
            info = dm.extract_requirement_info(conversation_history)
            is_complete = dm.check_info_complete(info)
            
            if not is_complete:
                if not info.get('requirement'):
                    reply = "请描述您的产品需求，我会帮您进行拆解和工时评估。"
                else:
                    reply = dm.get_next_question(info) or "请继续描述您的需求内容。"
                
                collect_data = {'type': 'complete', 'stage': 'collecting', 'message': reply, 'collected_info': info}
                yield 'data: ' + json.dumps(collect_data) + '\n\n'
                return
            
            # 拼装完整需求文本
            full_text = "【模块】%s 【类型】%s 【描述】%s" % (info['module'], info['type'], info['requirement'])
            
            # 构建历史上下文
            chat_keywords = ['你好', 'hello', 'hi', '您好', '早上好', '晚上好', '下午好', '嗨',
                            '谢谢', '谢谢了', '感谢', '拜拜', '再见', '再见了', '回见']
            msgs_ctx = []
            for m in conversation_history[:-1]:
                content = m['content'].lower()
                if not any(keyword in content for keyword in chat_keywords):
                    msgs_ctx.append(m)
            
            max_total_length = 3000
            history_ctx = ""
            for m in reversed(msgs_ctx):
                msg_text = "%s: %s" % (m['role'], m['content'])
                if len(history_ctx) + len(msg_text) <= max_total_length:
                    history_ctx = msg_text + "\n" + history_ctx
                else:
                    break
            
            # 发送thinking状态
            thinking_steps = [
                "正在分析需求...",
                "正在加载知识库...",
                "正在识别需求类型...",
                "正在进行结构化拆解...",
                "正在核算工时...",
                "正在整理结果...",
            ]
            
            for step in thinking_steps:
                thinking_data = {'type': 'thinking', 'message': step}
                yield 'data: ' + json.dumps(thinking_data) + '\n\n'
                time.sleep(0.3)
            
            # 执行评估
            try:
                session_knowledge = session.get_temp_knowledge_context()
                result = worktime_agent.run_chat(
                    text=full_text,
                    model_id=model_id,
                    context=history_ctx,
                    skill_id=skill_id,
                    last_evaluation=last_evaluation,
                    session_knowledge=session_knowledge,
                )
            except Exception as e:
                error_data = {'type': 'error', 'message': '评估失败: ' + str(e)}
                yield 'data: ' + json.dumps(error_data) + '\n\n'
                return
            
            if result.get("needs_clarification"):
                clarify_data = {'type': 'complete', 'stage': 'clarifying', 'message': result['clarification_question']}
                yield 'data: ' + json.dumps(clarify_data) + '\n\n'
                return
            
            # 保存评估结果
            session.set_last_evaluation({
                "total_days": result["total_days"],
                "role_breakdown": result.get("role_breakdown", {}),
                "g_text": result["g_text"],
                "pages_features": result.get("pages_features", []),
            })

            # 生成表格格式的输出
            pages_features = result.get("pages_features", [])
            table_output = worktime_agent.format_evaluation_as_table({
                "pages_features": pages_features,
                "total_days": result["total_days"],
                "role_breakdown": result.get("role_breakdown", {}),
            }, session_knowledge=session_knowledge)

            session.add_message("assistant", table_output)

            # 发送评估结果
            result_data = {
                'type': 'complete',
                'stage': 'assessment',
                'intent': intent,
                'intent_reason': intent_reason,
                'g_text': table_output,
                'total_days': result['total_days'],
                'role_breakdown': result.get('role_breakdown', {}),
                'pages_features': result.get('pages_features', []),
            }
            yield 'data: ' + json.dumps(result_data) + '\n\n'
            return

    return Response(generate(), mimetype='text/event-stream')



@app.route("/download/<path:filename>")
def download(filename):
    # 1. 安全过滤文件名（移除路径字符）
    safe_filename = secure_filename(filename)
    
    # 2. 构建完整路径
    filepath = os.path.join(OUTPUT_FOLDER, safe_filename)
    
    # 3. 强制校验：确保路径在允许范围内
    abs_output_folder = os.path.abspath(OUTPUT_FOLDER)
    abs_filepath = os.path.abspath(filepath)
    
    if not abs_filepath.startswith(abs_output_folder):
        logger.warning(f"非法路径访问尝试: {filename}")
        return jsonify({"error": "无效的文件请求"}), 403
    
    # 4. 检查文件是否存在
    if not os.path.exists(filepath):
        logger.warning(f"下载文件不存在: {filename}")
        return jsonify({"error": "文件不存在"}), 404
    
    logger.info(f"文件下载: {safe_filename}")
    return send_file(filepath, as_attachment=True, download_name=safe_filename)


# ============================================================

# ============================================================
# 技能管理接口
# ============================================================

@app.route("/skills", methods=["GET"])
def list_skills():
    """列出所有可用技能"""
    from agent import skill_manager as sm
    skills = sm.list_skills()
    current = sm.get_current_skill_id()
    return jsonify({"skills": skills, "current": current})


@app.route("/skills/current", methods=["GET"])
def get_current_skill():
    """获取当前技能详情"""
    from agent import skill_manager as sm
    sid   = sm.get_current_skill_id()
    skill = sm.get_skill(sid)
    return jsonify({"skill_id": sid, "skill": skill})


@app.route("/skills/switch", methods=["POST"])
def switch_skill():
    """切换当前全局技能"""
    from agent import skill_manager as sm
    data     = request.get_json()
    skill_id = (data.get("skill_id") or "").strip()
    if not skill_id:
        return jsonify({"error": "skill_id 不能为空"}), 400
    ok = sm.set_current_skill(skill_id)
    if not ok:
        available = [s["id"] for s in sm.list_skills()]
        return jsonify({"error": f"技能不存在: {skill_id}", "available": available}), 404
    logger.info(f"技能切换: {skill_id}")
    return jsonify({"success": True, "current": skill_id})


@app.route("/skills/<skill_id>", methods=["GET"])
def get_skill(skill_id):
    """获取指定技能的完整配置"""
    from agent import skill_manager as sm
    skill = sm.get_skill(skill_id)
    return jsonify({"skill_id": skill_id, "skill": skill})


@app.route("/skills/reload", methods=["POST"])
def reload_skill():
    """清除技能缓存，强制从文件重新加载（修改 JSON 后调用）"""
    from agent import skill_manager as sm
    data     = request.get_json() or {}
    skill_id = data.get("skill_id")
    if skill_id:
        sm.reload_skill(skill_id)
        logger.info(f"技能缓存已清除: {skill_id}")
        return jsonify({"success": True, "reloaded": skill_id})
    # 清除所有缓存
    sm._skills_cache.clear()
    logger.info("所有技能缓存已清除")
    return jsonify({"success": True, "reloaded": "all"})


@app.route("/skills/<skill_id>/examples", methods=["GET"])
def get_skill_examples(skill_id):
    """获取指定技能的历史评估案例"""
    from agent import skill_manager as sm
    limit    = int(request.args.get("limit", 20))
    examples = sm.load_examples(skill_id, limit=limit)
    return jsonify({"skill_id": skill_id, "count": len(examples), "examples": examples})


@app.route("/skills/<skill_id>/examples", methods=["POST"])
def add_skill_example(skill_id):
    """
    添加一条历史评估案例（用于学习）
    Body: {
      "requirement": {"module": "", "feature": "", "detail": ""},
      "pages_features": [...],
      "worktime": {"role_breakdown": {...}, "total_days": 0, "actual_days": 0, "note": ""}
    }
    """
    from agent import skill_manager as sm
    example = request.get_json()
    if not example or not example.get("requirement"):
        return jsonify({"error": "缺少 requirement 字段"}), 400
    fpath = sm.add_example(example, skill_id=skill_id)
    logger.info(f"新增历史案例: skill={skill_id} file={fpath}")
    return jsonify({"success": True, "file": os.path.basename(fpath)})


@app.route("/knowledge/code", methods=["GET"])
def list_code_knowledge():
    """列出代码知识库中的文件"""
    from agent import skill_manager as sm
    code_dir = sm.CODE_KB_DIR
    if not os.path.exists(code_dir):
        return jsonify({"files": [], "hint": "目录不存在，请创建 knowledge/code_knowledge/ 并放入文档"})
    files = [
        {"name": f, "size": os.path.getsize(os.path.join(code_dir, f))}
        for f in sorted(os.listdir(code_dir))
        if os.path.isfile(os.path.join(code_dir, f)) and not f.startswith(".")
    ]
    return jsonify({"files": files, "directory": code_dir})


@app.route("/export_evaluation", methods=["POST"])
def export_evaluation():
    """
    导出评估结果到Excel
    :param results: 评估结果数据
    :return: 文件下载路径
    """
    data = request.get_json()
    results = data.get("results")
    
    if not results:
        return jsonify({"error": "评估结果为空"}), 400
    
    logger.info(f"导出请求: count={len(results.get('results', []))}")
    
    try:
        # 生成导出文件路径
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        output_filename = f"evaluation_result_{timestamp}.xlsx"
        output_path = os.path.join(OUTPUT_FOLDER, output_filename)
        
        worktime_agent.export_to_excel(results, output_path)
        
        return jsonify({
            "success": True,
            "filename": output_filename,
            "download_url": f"/download/{output_filename}",
            "total_days": results.get("total_days", 0),
        })
    except Exception as e:
        logger.error(f"导出失败: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/knowledge/reload", methods=["POST"])
def reload_knowledge():
    """重新加载知识库（热更新）"""
    from agent.knowledge_manager import get_knowledge_manager
    
    kb_manager = get_knowledge_manager()
    kb_manager.kb_cache = {}
    kb_manager.code_kb_cache = {}
    
    try:
        all_kb = kb_manager.load_all_knowledge()
        logger.info("知识库重新加载成功")
        return jsonify({
            "success": True,
            "message": "知识库重新加载成功",
            "system_caps_count": len(all_kb.get("system_caps", {})),
            "business_docs_count": len(all_kb.get("business_docs", [])),
            "code_modules_count": len(all_kb.get("code_knowledge", {}).get("modules", [])),
        })
    except Exception as e:
        logger.error(f"知识库加载失败: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/knowledge/analyze", methods=["POST"])
def analyze_requirement():
    """
    分析需求（判断新增/调整，识别相关模块）
    :param text: 需求文本
    :return: 分析结果
    """
    data = request.get_json()
    text = (data.get("text") or "").strip()
    
    if not text:
        return jsonify({"error": "需求内容为空"}), 400
    
    logger.info(f"需求分析请求: text_length={len(text)}")
    
    try:
        from agent.knowledge_manager import get_knowledge_manager
        
        kb_manager = get_knowledge_manager()
        
        lines = text.split('\n')
        feature = lines[0].strip()[:80]
        detail_lines = [l for l in lines[1:] if l.strip()]
        detail = '\n'.join(detail_lines) if detail_lines else text
        
        req = {
            "module": "",
            "feature": feature,
            "detail": detail,
        }
        
        analysis = kb_manager.analyze_requirement(req)
        decomposition = kb_manager.suggest_decomposition(req)
        
        return jsonify({
            "success": True,
            "analysis": analysis,
            "decomposition": decomposition,
        })
    except Exception as e:
        logger.error(f"需求分析失败: {e}")
        return jsonify({"error": str(e)}), 500


if __name__ == "__main__":
    logger.info("启动 AI 工时评估助手服务")
    app.run(host="0.0.0.0", port=5001, debug=True)
